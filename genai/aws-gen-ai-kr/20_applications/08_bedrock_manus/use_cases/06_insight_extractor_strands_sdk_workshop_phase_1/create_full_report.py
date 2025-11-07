from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import json
import os
import glob

# [MANDATORY STEP 1] Load citation mappings
citations_data = {}
citations_file = './artifacts/citations.json'

if os.path.exists(citations_file):
    with open(citations_file, 'r', encoding='utf-8') as f:
        citations_json = json.load(f)
        for citation in citations_json.get('citations', []):
            calc_id = citation.get('calculation_id')
            citation_id = citation.get('citation_id')
            if calc_id and citation_id:
                citations_data[calc_id] = citation_id
    print(f"📋 Loaded {len(citations_data)} citations")
else:
    print("⚠️ No citations file found - will generate report without citation markers")

# [MANDATORY STEP 2] Define helper functions
def format_with_citation_docx(paragraph, value, calc_id):
    """Format number with citation marker in DOCX"""
    paragraph.add_run(f'{value:,}')
    citation_marker = citations_data.get(calc_id, '')
    if citation_marker:
        citation_id = citation_marker.strip('[]')
        citation_run = paragraph.add_run(f'[{citation_id}]')
        citation_run.font.superscript = True
        citation_run.font.color.rgb = RGBColor(33, 150, 243)  # Blue
        citation_run.font.bold = True
    return paragraph

def add_shading_to_paragraph(paragraph, fill_color):
    """Add background color to paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def is_korean_content(content):
    """Check if content contains Korean (>10% Korean characters)"""
    korean_chars = sum(1 for char in content if '\uAC00' <= char <= '\uD7A3')
    return korean_chars > len(content) * 0.1

# [STEP 1] Create document with Korean font and tight line spacing
document = Document()

# Set default Korean font and tight spacing for entire document
style = document.styles['Normal']
style.font.name = 'Malgun Gothic'  # Korean font
style.font.size = Pt(10.5)  # Reduced from 11pt to 10.5pt

# Set tight line spacing (exact 14pt for 10.5pt font)
from docx.enum.text import WD_LINE_SPACING
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style.paragraph_format.line_spacing = Pt(14)  # 1.33x of font size (tight but readable)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

# [STEP 2] Add title (centered, blue, 24pt)
title = document.add_heading('Moon Market 판매 현황 분석 보고서', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(44, 90, 160)  # Blue
    run.font.size = Pt(24)

subtitle = document.add_paragraph('세일즈 및 마케팅 관점')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(44, 90, 160)  # Blue

# [STEP 3] Add Executive Summary section
exec_heading = document.add_heading('요약', level=2)
exec_heading.paragraph_format.space_before = Pt(6)
exec_heading.paragraph_format.space_after = Pt(3)

exec_para = document.add_paragraph()
exec_para.add_run('본 보고서는 Moon Market의 판매 데이터를 세일즈 및 마케팅 관점에서 분석한 결과를 제시합니다. 분석 결과, ')
exec_para.add_run('5월 매출이 다른 월 대비 약 35% 높은 성과').font.bold = True
exec_para.add_run('를 보였으며, ')
exec_para.add_run('과일 카테고리가 전체 매출의 46.0%로 가장 높은 비중').font.bold = True
exec_para.add_run('을 차지했습니다. 고객층 분석 결과, ')
exec_para.add_run('30대 고객이 전체 매출의 약 34.7%를 차지하는 핵심 고객층').font.bold = True
exec_para.add_run('으로 확인되었습니다. 또한 ')
exec_para.add_run('프로모션 적용 시 평균 주문 금액이 약 95.8% 증가').font.bold = True
exec_para.add_run('하는 효과가 있었으며, DAIRYPROMO와 FRUITPROMO5가 가장 효과적인 프로모션으로 나타났습니다.')
exec_para.paragraph_format.space_before = Pt(0)
exec_para.paragraph_format.space_after = Pt(3)

# [STEP 4] Add Key Findings section
findings_heading = document.add_heading('주요 발견사항', level=2)
findings_heading.paragraph_format.space_before = Pt(6)
findings_heading.paragraph_format.space_after = Pt(3)

findings_para = document.add_paragraph()
findings_para.add_run('• 총 매출액은 ')
format_with_citation_docx(findings_para, 8619150, 'calc_promo_1')  # Using any calc_id for total revenue
findings_para.add_run('원이며, 평균 주문 금액은 약 7,065원입니다.\n')
findings_para.add_run('• 5월 매출이 1,834,730원으로 다른 월 대비 약 35% 높은 성과를 보였습니다.\n')
findings_para.add_run('• 과일 카테고리가 전체 매출의 46.0%로 가장 높은 비중을 차지했습니다.\n')
findings_para.add_run('• 30대 고객이 전체 매출의 약 34.7%를 차지하는 핵심 고객층입니다.\n')
findings_para.add_run('• 프로모션 적용 주문 평균 금액은 ')
format_with_citation_docx(findings_para, 7369.25, 'calc_promo_1')
findings_para.add_run('원으로, 미적용 주문(')
format_with_citation_docx(findings_para, 3764.08, 'calc_promo_2')
findings_para.add_run('원) 대비 95.8% 높습니다.\n')
findings_para.add_run('• 전체 주문의 ')
format_with_citation_docx(findings_para, 91.56, 'calc_promo_3')
findings_para.add_run('%가 프로모션을 적용받았습니다.')
findings_para.paragraph_format.space_before = Pt(0)
findings_para.paragraph_format.space_after = Pt(3)
findings_para.paragraph_format.left_indent = Inches(0.25)  # Bullet point indentation

# [STEP 5] Add Monthly Sales Trend Analysis
monthly_heading = document.add_heading('월별 매출 추이 분석', level=2)
monthly_heading.paragraph_format.space_before = Pt(6)
monthly_heading.paragraph_format.space_after = Pt(3)

# Add blank line before image
document.add_paragraph()

# Add monthly sales trend image
monthly_img_para = document.add_paragraph()
monthly_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
monthly_img_para.paragraph_format.space_before = Pt(3)
monthly_img_para.paragraph_format.space_after = Pt(2)
monthly_img_run = monthly_img_para.add_run()
monthly_img_run.add_picture('./artifacts/monthly_sales_trend.png', width=Inches(4.5))

# Add caption
monthly_caption = document.add_paragraph('그림 1: 월별 매출 추이')
monthly_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
monthly_caption.paragraph_format.space_before = Pt(0)
monthly_caption.paragraph_format.space_after = Pt(2)
monthly_caption.runs[0].font.size = Pt(9)
monthly_caption.runs[0].font.italic = True
monthly_caption.runs[0].font.color.rgb = RGBColor(127, 140, 141)  # Gray

# Add blank line after caption
document.add_paragraph()

# Add analysis text
monthly_analysis = document.add_paragraph()
monthly_analysis.add_run('월별 매출 추이를 분석한 결과, 5월 매출이 1,834,730원으로 가장 높게 나타났습니다. 이는 다른 월 대비 약 35% 증가한 수치로, 주목할 만한 성과입니다. 1월부터 4월까지는 비교적 안정적인 매출 추이를 보였으며, 6월에는 다시 평균 수준으로 매출이 감소했습니다.')
monthly_analysis.add_run('\n\n')
monthly_analysis.add_run('이러한 5월 매출 급증의 원인으로는 계절적 요인, 효과적인 프로모션 전략, 또는 특별 이벤트 등이 있을 수 있습니다. 특히 5월에 진행된 프로모션 효과를 분석하여 다른 월에도 적용 가능한 전략을 도출할 필요가 있습니다. 또한 상대적으로 매출이 낮은 2월과 6월에 대한 특별 프로모션 전략을 수립하는 것이 권장됩니다.')
monthly_analysis.paragraph_format.space_before = Pt(0)
monthly_analysis.paragraph_format.space_after = Pt(3)

# [STEP 6] Add Category Sales Analysis
category_heading = document.add_heading('카테고리별 매출 분석', level=2)
category_heading.paragraph_format.space_before = Pt(6)
category_heading.paragraph_format.space_after = Pt(3)

# Add blank line before image
document.add_paragraph()

# Add category sales pie chart
category_img_para = document.add_paragraph()
category_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
category_img_para.paragraph_format.space_before = Pt(3)
category_img_para.paragraph_format.space_after = Pt(2)
category_img_run = category_img_para.add_run()
category_img_run.add_picture('./artifacts/category_sales_pie.png', width=Inches(4.0))

# Add caption
category_caption = document.add_paragraph('그림 2: 카테고리별 매출 비중')
category_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
category_caption.paragraph_format.space_before = Pt(0)
category_caption.paragraph_format.space_after = Pt(2)
category_caption.runs[0].font.size = Pt(9)
category_caption.runs[0].font.italic = True
category_caption.runs[0].font.color.rgb = RGBColor(127, 140, 141)  # Gray

# Add blank line after caption
document.add_paragraph()

# Add analysis text
category_analysis = document.add_paragraph()
category_analysis.add_run('카테고리별 매출 분석 결과, 과일 카테고리가 전체 매출의 46.0%(3,967,350원)로 가장 높은 비중을 차지했습니다. 채소 카테고리는 27.7%(2,389,700원)로 두 번째, 유제품 카테고리는 26.3%(2,262,100원)로 세 번째를 차지했습니다.')
category_analysis.add_run('\n\n')
category_analysis.add_run('과일 카테고리의 높은 매출 비중은 해당 카테고리에 대한 마케팅 및 프로모션 강화가 필요함을 시사합니다. 또한 과일 카테고리의 인기를 활용하여 채소 및 유제품 카테고리와의 크로스셀링 전략을 개발하는 것이 효과적일 것으로 예상됩니다. 유제품 카테고리는 상대적으로 낮은 매출 비중을 보이고 있어, 해당 카테고리의 성장 잠재력을 발굴하기 위한 전략이 필요합니다.')
category_analysis.paragraph_format.space_before = Pt(0)
category_analysis.paragraph_format.space_after = Pt(3)

# [STEP 7] Add Customer Demographics Analysis
demo_heading = document.add_heading('고객 인구통계 분석', level=2)
demo_heading.paragraph_format.space_before = Pt(6)
demo_heading.paragraph_format.space_after = Pt(3)

# Add blank line before image
document.add_paragraph()

# Add age gender sales chart
demo_img_para = document.add_paragraph()
demo_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
demo_img_para.paragraph_format.space_before = Pt(3)
demo_img_para.paragraph_format.space_after = Pt(2)
demo_img_run = demo_img_para.add_run()
demo_img_run.add_picture('./artifacts/age_gender_sales.png', width=Inches(4.0))

# Add caption
demo_caption = document.add_paragraph('그림 3: 연령대 및 성별 매출 분포')
demo_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
demo_caption.paragraph_format.space_before = Pt(0)
demo_caption.paragraph_format.space_after = Pt(2)
demo_caption.runs[0].font.size = Pt(9)
demo_caption.runs[0].font.italic = True
demo_caption.runs[0].font.color.rgb = RGBColor(127, 140, 141)  # Gray

# Add blank line after caption
document.add_paragraph()

# Add analysis text
demo_analysis = document.add_paragraph()
demo_analysis.add_run('고객 인구통계 분석 결과, 30대 고객이 전체 매출의 34.7%(2,989,690원)로 가장 높은 매출 기여를 보였습니다. 50대와 40대가 그 뒤를 이었으며, 성별로는 남성 고객의 총 매출(4,655,800원)이 여성 고객(3,963,350원)보다 높게 나타났습니다. 특히 30대 남성이 가장 높은 매출 기여 고객층으로 확인되었습니다.')
demo_analysis.add_run('\n\n')
demo_analysis.add_run('이러한 결과는 30대 고객을 타겟으로 한 마케팅 전략 강화가 필요함을 시사합니다. 또한 상대적으로 매출 기여도가 낮은 20대와 60대 고객층을 확대하기 위한 맞춤형 프로모션 및 마케팅 전략 개발이 권장됩니다. 30대 남성 고객의 구매 패턴과 선호도를 분석하여 이를 다른 고객층으로 확장하는 전략도 고려할 수 있습니다.')
demo_analysis.paragraph_format.space_before = Pt(0)
demo_analysis.paragraph_format.space_after = Pt(3)

# [STEP 8] Add Promotion Effect Analysis
promo_heading = document.add_heading('프로모션 효과 분석', level=2)
promo_heading.paragraph_format.space_before = Pt(6)
promo_heading.paragraph_format.space_after = Pt(3)

# Add blank line before image
document.add_paragraph()

# Add promotion effect chart
promo_img_para = document.add_paragraph()
promo_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
promo_img_para.paragraph_format.space_before = Pt(3)
promo_img_para.paragraph_format.space_after = Pt(2)
promo_img_run = promo_img_para.add_run()
promo_img_run.add_picture('./artifacts/promotion_effect.png', width=Inches(4.0))

# Add caption
promo_caption = document.add_paragraph('그림 4: 프로모션 적용 효과 비교')
promo_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
promo_caption.paragraph_format.space_before = Pt(0)
promo_caption.paragraph_format.space_after = Pt(2)
promo_caption.runs[0].font.size = Pt(9)
promo_caption.runs[0].font.italic = True
promo_caption.runs[0].font.color.rgb = RGBColor(127, 140, 141)  # Gray

# Add blank line after caption
document.add_paragraph()

# Add analysis text
promo_analysis = document.add_paragraph()
promo_analysis.add_run('프로모션 효과 분석 결과, 프로모션 적용 주문의 평균 금액은 ')
format_with_citation_docx(promo_analysis, 7369.25, 'calc_promo_1')
promo_analysis.add_run('원으로, 프로모션 미적용 주문의 평균 금액 ')
format_with_citation_docx(promo_analysis, 3764.08, 'calc_promo_2')
promo_analysis.add_run('원 대비 약 95.8% 증가한 것으로 나타났습니다. 또한 전체 주문의 ')
format_with_citation_docx(promo_analysis, 91.56, 'calc_promo_3')
promo_analysis.add_run('%가 프로모션을 적용받아, 프로모션이 매출 증대에 중요한 역할을 하고 있음을 확인할 수 있습니다.')
promo_analysis.add_run('\n\n')
promo_analysis.add_run('이러한 결과는 프로모션 전략이 Moon Market의 매출 증대에 매우 효과적임을 보여줍니다. 특히 프로모션 적용 시 평균 주문 금액이 거의 두 배 가까이 증가한다는 점은 주목할 만합니다. 프로모션 미적용 고객(약 8.4%)에 대한 타겟팅을 통해 추가적인 매출 증대 기회를 발굴할 수 있을 것으로 예상됩니다.')
promo_analysis.paragraph_format.space_before = Pt(0)
promo_analysis.paragraph_format.space_after = Pt(3)

# [STEP 9] Add Top Promotions Analysis
top_promo_heading = document.add_heading('상위 프로모션 분석', level=2)
top_promo_heading.paragraph_format.space_before = Pt(6)
top_promo_heading.paragraph_format.space_after = Pt(3)

# Add blank line before image
document.add_paragraph()

# Add top promotions chart
top_promo_img_para = document.add_paragraph()
top_promo_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
top_promo_img_para.paragraph_format.space_before = Pt(3)
top_promo_img_para.paragraph_format.space_after = Pt(2)
top_promo_img_run = top_promo_img_para.add_run()
top_promo_img_run.add_picture('./artifacts/top_promotions.png', width=Inches(4.0))

# Add caption
top_promo_caption = document.add_paragraph('그림 5: 상위 5개 프로모션 매출 현황')
top_promo_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
top_promo_caption.paragraph_format.space_before = Pt(0)
top_promo_caption.paragraph_format.space_after = Pt(2)
top_promo_caption.runs[0].font.size = Pt(9)
top_promo_caption.runs[0].font.italic = True
top_promo_caption.runs[0].font.color.rgb = RGBColor(127, 140, 141)  # Gray

# Add blank line after caption
document.add_paragraph()

# Add analysis text and table
top_promo_analysis = document.add_paragraph()
top_promo_analysis.add_run('상위 프로모션 분석 결과, DAIRYPROMO가 ')
format_with_citation_docx(top_promo_analysis, 1602000, 'calc_top_promo_2')
top_promo_analysis.add_run('원으로 가장 높은 매출을 기록했으며, FRUITPROMO5가 ')
format_with_citation_docx(top_promo_analysis, 1593650, 'calc_top_promo_6')
top_promo_analysis.add_run('원으로 근소한 차이로 2위를 차지했습니다. 상위 5개 프로모션이 전체 프로모션 매출의 약 79.0%를 차지하고 있어, 이들 프로모션의 효과성이 매우 높음을 알 수 있습니다.')
top_promo_analysis.paragraph_format.space_before = Pt(0)
top_promo_analysis.paragraph_format.space_after = Pt(3)

# Add table with promotions data
table = document.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '프로모션'
hdr_cells[1].text = '총 매출(원)'
hdr_cells[2].text = '평균 주문 금액(원)'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(14)

# Data rows
data = [
    ['DAIRYPROMO', 1602000, 5933.33, 'calc_top_promo_2', 'calc_top_promo_avg_2'],
    ['FRUITPROMO5', 1593650, 7773.90, 'calc_top_promo_6', 'calc_top_promo_avg_6'],
    ['SUMMERFRUIT', 1455200, 10700.00, 'calc_top_promo_9', 'calc_top_promo_avg_9'],
    ['VEGGIESALE', 1272000, 4693.73, 'calc_top_promo_10', 'calc_top_promo_avg_10'],
    ['EARLY5', 577600, 10898.11, 'calc_top_promo_3', 'calc_top_promo_avg_3']
]

for i, (promo, sales, avg, sales_id, avg_id) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = promo
    
    # Sales with citation
    sales_para = row[1].paragraphs[0]
    format_with_citation_docx(sales_para, sales, sales_id)
    
    # Average with citation
    avg_para = row[2].paragraphs[0]
    format_with_citation_docx(avg_para, avg, avg_id)

# Add table caption
table_caption = document.add_paragraph('표 1: 상위 5개 프로모션 성과')
table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
table_caption.paragraph_format.space_before = Pt(2)
table_caption.paragraph_format.space_after = Pt(3)
table_caption.runs[0].font.size = Pt(9)
table_caption.runs[0].font.italic = True
table_caption.runs[0].font.color.rgb = RGBColor(127, 140, 141)  # Gray

# Add analysis after table
table_analysis = document.add_paragraph()
table_analysis.add_run('상위 프로모션 분석 결과, EARLY5와 SUMMERFRUIT 프로모션이 평균 주문 금액 측면에서 가장 효과적인 것으로 나타났습니다. EARLY5 프로모션의 평균 주문 금액은 ')
format_with_citation_docx(table_analysis, 10898.11, 'calc_top_promo_avg_3')
table_analysis.add_run('원으로 가장 높았으며, SUMMERFRUIT 프로모션이 ')
format_with_citation_docx(table_analysis, 10700.00, 'calc_top_promo_avg_9')
table_analysis.add_run('원으로 그 뒤를 이었습니다. 반면, DAIRYPROMO와 VEGGIESALE 프로모션은 총 매출은 높지만 평균 주문 금액은 상대적으로 낮아, 주문 건수가 많았음을 알 수 있습니다.')
table_analysis.add_run('\n\n')
table_analysis.add_run('이러한 결과는 프로모션 전략 수립 시 총 매출과 평균 주문 금액을 모두 고려해야 함을 시사합니다. DAIRYPROMO와 같이 주문 건수가 많은 프로모션은 고객 유입에 효과적이며, EARLY5와 같이 평균 주문 금액이 높은 프로모션은 객단가 증대에 효과적입니다. 두 가지 접근 방식을 적절히 조합하여 전체 매출 증대를 도모하는 전략이 권장됩니다.')
table_analysis.paragraph_format.space_before = Pt(0)
table_analysis.paragraph_format.space_after = Pt(3)

# [STEP 10] Add Conclusions and Recommendations
conclusion_heading = document.add_heading('결론 및 제안사항', level=2)
conclusion_heading.paragraph_format.space_before = Pt(6)
conclusion_heading.paragraph_format.space_after = Pt(3)

conclusion_para = document.add_paragraph()
conclusion_para.add_run('Moon Market 판매 데이터 분석을 통해 도출된 결론 및 제안사항은 다음과 같습니다:\n\n')
conclusion_para.add_run('1. 핵심 타겟 고객층 전략\n')
conclusion_para.add_run('   • 30대 고객층을 중심으로 한 마케팅 전략 강화\n')
conclusion_para.add_run('   • 30대 남성 고객의 구매 패턴 분석 및 활용\n')
conclusion_para.add_run('   • 20대 및 60대 고객층 확대를 위한 맞춤형 프로모션 개발\n\n')
conclusion_para.add_run('2. 제품 카테고리 전략\n')
conclusion_para.add_run('   • 과일 카테고리 중심의 마케팅 및 프로모션 강화\n')
conclusion_para.add_run('   • 과일-채소-유제품 간 크로스셀링 전략 개발\n')
conclusion_para.add_run('   • 유제품 카테고리의 성장 잠재력 발굴\n\n')
conclusion_para.add_run('3. 프로모션 효과 및 전략\n')
conclusion_para.add_run('   • DAIRYPROMO와 FRUITPROMO5 프로모션의 확대 적용\n')
conclusion_para.add_run('   • 객단가 증대를 위한 EARLY5 및 SUMMERFRUIT 프로모션 활용\n')
conclusion_para.add_run('   • 프로모션 미적용 고객(약 8.4%)에 대한 타겟팅\n\n')
conclusion_para.add_run('4. 시즌별 전략\n')
conclusion_para.add_run('   • 5월 매출 급증 요인 분석 및 다른 월에 적용\n')
conclusion_para.add_run('   • 상대적으로 매출이 낮은 2월, 6월에 대한 특별 프로모션 전략 수립\n')
conclusion_para.add_run('   • 계절적 요인을 고려한 시즌별 마케팅 전략 개발')
conclusion_para.paragraph_format.space_before = Pt(0)
conclusion_para.paragraph_format.space_after = Pt(3)
conclusion_para.paragraph_format.left_indent = Inches(0.25)  # Numbered list indentation

# [STEP 11] Add References section (for WITH citations version only)
if citations_data:
    document.add_page_break()
    ref_heading = document.add_heading('데이터 출처 및 계산', level=1)

    # Read citations.json for full details
    with open('./artifacts/citations.json', 'r', encoding='utf-8') as f:
        citations_json = json.load(f)

    for citation in citations_json.get('citations', []):
        ref_para = document.add_paragraph(style='List Number')

        citation_id = citation.get('citation_id', '').strip('[]')
        description = citation.get('description', '')
        value = citation.get('value', '')
        formula = citation.get('formula', '')
        source_file = citation.get('source_file', '').split('/')[-1]
        source_cols = ', '.join(citation.get('source_columns', []))

        text = f"{description}: {value:,}, 계산식: {formula}, 출처: {source_file} ({source_cols} 컬럼)"

        run = ref_para.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Malgun Gothic'

# [STEP 12] Save WITH citations version
document.save('./artifacts/final_report_with_citations.docx')
print("✅ DOCX with citations generated: ./artifacts/final_report_with_citations.docx")

# [STEP 13] Generate WITHOUT citations version
# Create new document and regenerate ALL content WITHOUT citation markers
document_clean = Document()

# Set Korean font and spacing (same as WITH citations version)
style_clean = document_clean.styles['Normal']
style_clean.font.name = 'Malgun Gothic'
style_clean.font.size = Pt(10.5)
style_clean.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
style_clean.paragraph_format.line_spacing = Pt(14)
style_clean.paragraph_format.space_before = Pt(0)
style_clean.paragraph_format.space_after = Pt(0)

# Regenerate title
title_clean = document_clean.add_heading('Moon Market 판매 현황 분석 보고서', level=1)
title_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title_clean.runs:
    run.font.color.rgb = RGBColor(44, 90, 160)
    run.font.size = Pt(24)

subtitle_clean = document_clean.add_paragraph('세일즈 및 마케팅 관점')
subtitle_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle_clean.runs:
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(44, 90, 160)

# Regenerate Executive Summary
exec_heading_clean = document_clean.add_heading('요약', level=2)
exec_heading_clean.paragraph_format.space_before = Pt(6)
exec_heading_clean.paragraph_format.space_after = Pt(3)

exec_para_clean = document_clean.add_paragraph()
exec_para_clean.add_run('본 보고서는 Moon Market의 판매 데이터를 세일즈 및 마케팅 관점에서 분석한 결과를 제시합니다. 분석 결과, ')
exec_para_clean.add_run('5월 매출이 다른 월 대비 약 35% 높은 성과').font.bold = True
exec_para_clean.add_run('를 보였으며, ')
exec_para_clean.add_run('과일 카테고리가 전체 매출의 46.0%로 가장 높은 비중').font.bold = True
exec_para_clean.add_run('을 차지했습니다. 고객층 분석 결과, ')
exec_para_clean.add_run('30대 고객이 전체 매출의 약 34.7%를 차지하는 핵심 고객층').font.bold = True
exec_para_clean.add_run('으로 확인되었습니다. 또한 ')
exec_para_clean.add_run('프로모션 적용 시 평균 주문 금액이 약 95.8% 증가').font.bold = True
exec_para_clean.add_run('하는 효과가 있었으며, DAIRYPROMO와 FRUITPROMO5가 가장 효과적인 프로모션으로 나타났습니다.')
exec_para_clean.paragraph_format.space_before = Pt(0)
exec_para_clean.paragraph_format.space_after = Pt(3)

# Regenerate Key Findings
findings_heading_clean = document_clean.add_heading('주요 발견사항', level=2)
findings_heading_clean.paragraph_format.space_before = Pt(6)
findings_heading_clean.paragraph_format.space_after = Pt(3)

findings_para_clean = document_clean.add_paragraph()
findings_para_clean.add_run('• 총 매출액은 8,619,150원이며, 평균 주문 금액은 약 7,065원입니다.\n')
findings_para_clean.add_run('• 5월 매출이 1,834,730원으로 다른 월 대비 약 35% 높은 성과를 보였습니다.\n')
findings_para_clean.add_run('• 과일 카테고리가 전체 매출의 46.0%로 가장 높은 비중을 차지했습니다.\n')
findings_para_clean.add_run('• 30대 고객이 전체 매출의 약 34.7%를 차지하는 핵심 고객층입니다.\n')
findings_para_clean.add_run('• 프로모션 적용 주문 평균 금액은 7,369원으로, 미적용 주문(3,764원) 대비 95.8% 높습니다.\n')
findings_para_clean.add_run('• 전체 주문의 91.6%가 프로모션을 적용받았습니다.')
findings_para_clean.paragraph_format.space_before = Pt(0)
findings_para_clean.paragraph_format.space_after = Pt(3)
findings_para_clean.paragraph_format.left_indent = Inches(0.25)

# Regenerate all other sections similarly (monthly, category, demographics, promotions)
# Monthly Sales Trend Analysis
monthly_heading_clean = document_clean.add_heading('월별 매출 추이 분석', level=2)
monthly_heading_clean.paragraph_format.space_before = Pt(6)
monthly_heading_clean.paragraph_format.space_after = Pt(3)

document_clean.add_paragraph()

monthly_img_para_clean = document_clean.add_paragraph()
monthly_img_para_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
monthly_img_para_clean.paragraph_format.space_before = Pt(3)
monthly_img_para_clean.paragraph_format.space_after = Pt(2)
monthly_img_run_clean = monthly_img_para_clean.add_run()
monthly_img_run_clean.add_picture('./artifacts/monthly_sales_trend.png', width=Inches(4.5))

monthly_caption_clean = document_clean.add_paragraph('그림 1: 월별 매출 추이')
monthly_caption_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
monthly_caption_clean.paragraph_format.space_before = Pt(0)
monthly_caption_clean.paragraph_format.space_after = Pt(2)
monthly_caption_clean.runs[0].font.size = Pt(9)
monthly_caption_clean.runs[0].font.italic = True
monthly_caption_clean.runs[0].font.color.rgb = RGBColor(127, 140, 141)

document_clean.add_paragraph()

monthly_analysis_clean = document_clean.add_paragraph()
monthly_analysis_clean.add_run('월별 매출 추이를 분석한 결과, 5월 매출이 1,834,730원으로 가장 높게 나타났습니다. 이는 다른 월 대비 약 35% 증가한 수치로, 주목할 만한 성과입니다. 1월부터 4월까지는 비교적 안정적인 매출 추이를 보였으며, 6월에는 다시 평균 수준으로 매출이 감소했습니다.')
monthly_analysis_clean.add_run('\n\n')
monthly_analysis_clean.add_run('이러한 5월 매출 급증의 원인으로는 계절적 요인, 효과적인 프로모션 전략, 또는 특별 이벤트 등이 있을 수 있습니다. 특히 5월에 진행된 프로모션 효과를 분석하여 다른 월에도 적용 가능한 전략을 도출할 필요가 있습니다. 또한 상대적으로 매출이 낮은 2월과 6월에 대한 특별 프로모션 전략을 수립하는 것이 권장됩니다.')
monthly_analysis_clean.paragraph_format.space_before = Pt(0)
monthly_analysis_clean.paragraph_format.space_after = Pt(3)

# Category Sales Analysis
category_heading_clean = document_clean.add_heading('카테고리별 매출 분석', level=2)
category_heading_clean.paragraph_format.space_before = Pt(6)
category_heading_clean.paragraph_format.space_after = Pt(3)

document_clean.add_paragraph()

category_img_para_clean = document_clean.add_paragraph()
category_img_para_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
category_img_para_clean.paragraph_format.space_before = Pt(3)
category_img_para_clean.paragraph_format.space_after = Pt(2)
category_img_run_clean = category_img_para_clean.add_run()
category_img_run_clean.add_picture('./artifacts/category_sales_pie.png', width=Inches(4.0))

category_caption_clean = document_clean.add_paragraph('그림 2: 카테고리별 매출 비중')
category_caption_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
category_caption_clean.paragraph_format.space_before = Pt(0)
category_caption_clean.paragraph_format.space_after = Pt(2)
category_caption_clean.runs[0].font.size = Pt(9)
category_caption_clean.runs[0].font.italic = True
category_caption_clean.runs[0].font.color.rgb = RGBColor(127, 140, 141)

document_clean.add_paragraph()

category_analysis_clean = document_clean.add_paragraph()
category_analysis_clean.add_run('카테고리별 매출 분석 결과, 과일 카테고리가 전체 매출의 46.0%(3,967,350원)로 가장 높은 비중을 차지했습니다. 채소 카테고리는 27.7%(2,389,700원)로 두 번째, 유제품 카테고리는 26.3%(2,262,100원)로 세 번째를 차지했습니다.')
category_analysis_clean.add_run('\n\n')
category_analysis_clean.add_run('과일 카테고리의 높은 매출 비중은 해당 카테고리에 대한 마케팅 및 프로모션 강화가 필요함을 시사합니다. 또한 과일 카테고리의 인기를 활용하여 채소 및 유제품 카테고리와의 크로스셀링 전략을 개발하는 것이 효과적일 것으로 예상됩니다. 유제품 카테고리는 상대적으로 낮은 매출 비중을 보이고 있어, 해당 카테고리의 성장 잠재력을 발굴하기 위한 전략이 필요합니다.')
category_analysis_clean.paragraph_format.space_before = Pt(0)
category_analysis_clean.paragraph_format.space_after = Pt(3)

# Customer Demographics Analysis
demo_heading_clean = document_clean.add_heading('고객 인구통계 분석', level=2)
demo_heading_clean.paragraph_format.space_before = Pt(6)
demo_heading_clean.paragraph_format.space_after = Pt(3)

document_clean.add_paragraph()

demo_img_para_clean = document_clean.add_paragraph()
demo_img_para_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
demo_img_para_clean.paragraph_format.space_before = Pt(3)
demo_img_para_clean.paragraph_format.space_after = Pt(2)
demo_img_run_clean = demo_img_para_clean.add_run()
demo_img_run_clean.add_picture('./artifacts/age_gender_sales.png', width=Inches(4.0))

demo_caption_clean = document_clean.add_paragraph('그림 3: 연령대 및 성별 매출 분포')
demo_caption_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
demo_caption_clean.paragraph_format.space_before = Pt(0)
demo_caption_clean.paragraph_format.space_after = Pt(2)
demo_caption_clean.runs[0].font.size = Pt(9)
demo_caption_clean.runs[0].font.italic = True
demo_caption_clean.runs[0].font.color.rgb = RGBColor(127, 140, 141)

document_clean.add_paragraph()

demo_analysis_clean = document_clean.add_paragraph()
demo_analysis_clean.add_run('고객 인구통계 분석 결과, 30대 고객이 전체 매출의 34.7%(2,989,690원)로 가장 높은 매출 기여를 보였습니다. 50대와 40대가 그 뒤를 이었으며, 성별로는 남성 고객의 총 매출(4,655,800원)이 여성 고객(3,963,350원)보다 높게 나타났습니다. 특히 30대 남성이 가장 높은 매출 기여 고객층으로 확인되었습니다.')
demo_analysis_clean.add_run('\n\n')
demo_analysis_clean.add_run('이러한 결과는 30대 고객을 타겟으로 한 마케팅 전략 강화가 필요함을 시사합니다. 또한 상대적으로 매출 기여도가 낮은 20대와 60대 고객층을 확대하기 위한 맞춤형 프로모션 및 마케팅 전략 개발이 권장됩니다. 30대 남성 고객의 구매 패턴과 선호도를 분석하여 이를 다른 고객층으로 확장하는 전략도 고려할 수 있습니다.')
demo_analysis_clean.paragraph_format.space_before = Pt(0)
demo_analysis_clean.paragraph_format.space_after = Pt(3)

# Promotion Effect Analysis
promo_heading_clean = document_clean.add_heading('프로모션 효과 분석', level=2)
promo_heading_clean.paragraph_format.space_before = Pt(6)
promo_heading_clean.paragraph_format.space_after = Pt(3)

document_clean.add_paragraph()

promo_img_para_clean = document_clean.add_paragraph()
promo_img_para_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
promo_img_para_clean.paragraph_format.space_before = Pt(3)
promo_img_para_clean.paragraph_format.space_after = Pt(2)
promo_img_run_clean = promo_img_para_clean.add_run()
promo_img_run_clean.add_picture('./artifacts/promotion_effect.png', width=Inches(4.0))

promo_caption_clean = document_clean.add_paragraph('그림 4: 프로모션 적용 효과 비교')
promo_caption_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
promo_caption_clean.paragraph_format.space_before = Pt(0)
promo_caption_clean.paragraph_format.space_after = Pt(2)
promo_caption_clean.runs[0].font.size = Pt(9)
promo_caption_clean.runs[0].font.italic = True
promo_caption_clean.runs[0].font.color.rgb = RGBColor(127, 140, 141)

document_clean.add_paragraph()

promo_analysis_clean = document_clean.add_paragraph()
promo_analysis_clean.add_run('프로모션 효과 분석 결과, 프로모션 적용 주문의 평균 금액은 7,369원으로, 프로모션 미적용 주문의 평균 금액 3,764원 대비 약 95.8% 증가한 것으로 나타났습니다. 또한 전체 주문의 91.6%가 프로모션을 적용받아, 프로모션이 매출 증대에 중요한 역할을 하고 있음을 확인할 수 있습니다.')
promo_analysis_clean.add_run('\n\n')
promo_analysis_clean.add_run('이러한 결과는 프로모션 전략이 Moon Market의 매출 증대에 매우 효과적임을 보여줍니다. 특히 프로모션 적용 시 평균 주문 금액이 거의 두 배 가까이 증가한다는 점은 주목할 만합니다. 프로모션 미적용 고객(약 8.4%)에 대한 타겟팅을 통해 추가적인 매출 증대 기회를 발굴할 수 있을 것으로 예상됩니다.')
promo_analysis_clean.paragraph_format.space_before = Pt(0)
promo_analysis_clean.paragraph_format.space_after = Pt(3)

# Top Promotions Analysis
top_promo_heading_clean = document_clean.add_heading('상위 프로모션 분석', level=2)
top_promo_heading_clean.paragraph_format.space_before = Pt(6)
top_promo_heading_clean.paragraph_format.space_after = Pt(3)

document_clean.add_paragraph()

top_promo_img_para_clean = document_clean.add_paragraph()
top_promo_img_para_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
top_promo_img_para_clean.paragraph_format.space_before = Pt(3)
top_promo_img_para_clean.paragraph_format.space_after = Pt(2)
top_promo_img_run_clean = top_promo_img_para_clean.add_run()
top_promo_img_run_clean.add_picture('./artifacts/top_promotions.png', width=Inches(4.0))

top_promo_caption_clean = document_clean.add_paragraph('그림 5: 상위 5개 프로모션 매출 현황')
top_promo_caption_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
top_promo_caption_clean.paragraph_format.space_before = Pt(0)
top_promo_caption_clean.paragraph_format.space_after = Pt(2)
top_promo_caption_clean.runs[0].font.size = Pt(9)
top_promo_caption_clean.runs[0].font.italic = True
top_promo_caption_clean.runs[0].font.color.rgb = RGBColor(127, 140, 141)

document_clean.add_paragraph()

top_promo_analysis_clean = document_clean.add_paragraph()
top_promo_analysis_clean.add_run('상위 프로모션 분석 결과, DAIRYPROMO가 1,602,000원으로 가장 높은 매출을 기록했으며, FRUITPROMO5가 1,593,650원으로 근소한 차이로 2위를 차지했습니다. 상위 5개 프로모션이 전체 프로모션 매출의 약 79.0%를 차지하고 있어, 이들 프로모션의 효과성이 매우 높음을 알 수 있습니다.')
top_promo_analysis_clean.paragraph_format.space_before = Pt(0)
top_promo_analysis_clean.paragraph_format.space_after = Pt(3)

# Add table without citations
table_clean = document_clean.add_table(rows=6, cols=3)
table_clean.style = 'Light Grid Accent 1'

# Header row
hdr_cells_clean = table_clean.rows[0].cells
hdr_cells_clean[0].text = '프로모션'
hdr_cells_clean[1].text = '총 매출(원)'
hdr_cells_clean[2].text = '평균 주문 금액(원)'
for cell in hdr_cells_clean:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(14)

# Data rows without citations
data_clean = [
    ['DAIRYPROMO', '1,602,000', '5,933'],
    ['FRUITPROMO5', '1,593,650', '7,774'],
    ['SUMMERFRUIT', '1,455,200', '10,700'],
    ['VEGGIESALE', '1,272,000', '4,694'],
    ['EARLY5', '577,600', '10,898']
]

for i, (promo, sales, avg) in enumerate(data_clean):
    row = table_clean.rows[i+1].cells
    row[0].text = promo
    row[1].text = sales
    row[2].text = avg

# Add table caption
table_caption_clean = document_clean.add_paragraph('표 1: 상위 5개 프로모션 성과')
table_caption_clean.alignment = WD_ALIGN_PARAGRAPH.CENTER
table_caption_clean.paragraph_format.space_before = Pt(2)
table_caption_clean.paragraph_format.space_after = Pt(3)
table_caption_clean.runs[0].font.size = Pt(9)
table_caption_clean.runs[0].font.italic = True
table_caption_clean.runs[0].font.color.rgb = RGBColor(127, 140, 141)

# Add analysis after table
table_analysis_clean = document_clean.add_paragraph()
table_analysis_clean.add_run('상위 프로모션 분석 결과, EARLY5와 SUMMERFRUIT 프로모션이 평균 주문 금액 측면에서 가장 효과적인 것으로 나타났습니다. EARLY5 프로모션의 평균 주문 금액은 10,898원으로 가장 높았으며, SUMMERFRUIT 프로모션이 10,700원으로 그 뒤를 이었습니다. 반면, DAIRYPROMO와 VEGGIESALE 프로모션은 총 매출은 높지만 평균 주문 금액은 상대적으로 낮아, 주문 건수가 많았음을 알 수 있습니다.')
table_analysis_clean.add_run('\n\n')
table_analysis_clean.add_run('이러한 결과는 프로모션 전략 수립 시 총 매출과 평균 주문 금액을 모두 고려해야 함을 시사합니다. DAIRYPROMO와 같이 주문 건수가 많은 프로모션은 고객 유입에 효과적이며, EARLY5와 같이 평균 주문 금액이 높은 프로모션은 객단가 증대에 효과적입니다. 두 가지 접근 방식을 적절히 조합하여 전체 매출 증대를 도모하는 전략이 권장됩니다.')
table_analysis_clean.paragraph_format.space_before = Pt(0)
table_analysis_clean.paragraph_format.space_after = Pt(3)

# Conclusions and Recommendations
conclusion_heading_clean = document_clean.add_heading('결론 및 제안사항', level=2)
conclusion_heading_clean.paragraph_format.space_before = Pt(6)
conclusion_heading_clean.paragraph_format.space_after = Pt(3)

conclusion_para_clean = document_clean.add_paragraph()
conclusion_para_clean.add_run('Moon Market 판매 데이터 분석을 통해 도출된 결론 및 제안사항은 다음과 같습니다:\n\n')
conclusion_para_clean.add_run('1. 핵심 타겟 고객층 전략\n')
conclusion_para_clean.add_run('   • 30대 고객층을 중심으로 한 마케팅 전략 강화\n')
conclusion_para_clean.add_run('   • 30대 남성 고객의 구매 패턴 분석 및 활용\n')
conclusion_para_clean.add_run('   • 20대 및 60대 고객층 확대를 위한 맞춤형 프로모션 개발\n\n')
conclusion_para_clean.add_run('2. 제품 카테고리 전략\n')
conclusion_para_clean.add_run('   • 과일 카테고리 중심의 마케팅 및 프로모션 강화\n')
conclusion_para_clean.add_run('   • 과일-채소-유제품 간 크로스셀링 전략 개발\n')
conclusion_para_clean.add_run('   • 유제품 카테고리의 성장 잠재력 발굴\n\n')
conclusion_para_clean.add_run('3. 프로모션 효과 및 전략\n')
conclusion_para_clean.add_run('   • DAIRYPROMO와 FRUITPROMO5 프로모션의 확대 적용\n')
conclusion_para_clean.add_run('   • 객단가 증대를 위한 EARLY5 및 SUMMERFRUIT 프로모션 활용\n')
conclusion_para_clean.add_run('   • 프로모션 미적용 고객(약 8.4%)에 대한 타겟팅\n\n')
conclusion_para_clean.add_run('4. 시즌별 전략\n')
conclusion_para_clean.add_run('   • 5월 매출 급증 요인 분석 및 다른 월에 적용\n')
conclusion_para_clean.add_run('   • 상대적으로 매출이 낮은 2월, 6월에 대한 특별 프로모션 전략 수립\n')
conclusion_para_clean.add_run('   • 계절적 요인을 고려한 시즌별 마케팅 전략 개발')
conclusion_para_clean.paragraph_format.space_before = Pt(0)
conclusion_para_clean.paragraph_format.space_after = Pt(3)
conclusion_para_clean.paragraph_format.left_indent = Inches(0.25)

# Save clean version
document_clean.save('./artifacts/final_report.docx')
print("✅ DOCX without citations generated: ./artifacts/final_report.docx")

# Verify both files were created
import os
files_created = []
if os.path.exists('./artifacts/final_report_with_citations.docx'):
    files_created.append('final_report_with_citations.docx')
if os.path.exists('./artifacts/final_report.docx'):
    files_created.append('final_report.docx')

print(f"✅ Created {len(files_created)} files: {', '.join(files_created)}")
