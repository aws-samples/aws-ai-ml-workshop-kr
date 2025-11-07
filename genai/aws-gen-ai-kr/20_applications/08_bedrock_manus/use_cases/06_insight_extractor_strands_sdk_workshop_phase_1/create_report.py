from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document
document = Document()

# Add title
title = document.add_heading('Moon Market 판매 현황 분석 보고서', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add paragraph
p = document.add_paragraph('이 보고서는 Moon Market의 판매 데이터를 분석한 결과입니다.')

# Save document
document.save('./artifacts/final_report.docx')
print("Report created successfully")
