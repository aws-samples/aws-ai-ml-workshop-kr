"""
DOCX Setup Template for Reporter Agent

This file contains the mandatory setup code that must be executed FIRST
before any DOCX report generation to avoid NameError.

Usage in python_repl:
    exec(open('./src/prompts/templates/docx_setup_template.py').read())
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import json
import os
import glob

# [STEP 1] Load citation mappings
citations_data = {}
citations_file = './artifacts/citations.json'

if os.path.exists(citations_file):
    with open(citations_file, 'r', encoding='utf-8') as f:
        citations_json = json.load(f)
        for citation in citations_json.get('citations', []):
            calc_id = citation.get('calculation_id')
            citation_id = citation.get('citation_id')
            if calc_id and citation_id:
                citations_data[calc_id] = citation_id
    print(f"📋 Loaded {len(citations_data)} citations")
else:
    print("⚠️ No citations file found - will generate report without citation markers")

# [STEP 2] Define helper functions
def format_with_citation_docx(paragraph, value, calc_id):
    """Format number with citation marker in DOCX"""
    paragraph.add_run(f'{value:,}')
    citation_marker = citations_data.get(calc_id, '')
    if citation_marker:
        citation_id = citation_marker.strip('[]')
        citation_run = paragraph.add_run(f'[{citation_id}]')
        citation_run.font.superscript = True
        citation_run.font.color.rgb = RGBColor(33, 150, 243)  # Blue
        citation_run.font.bold = True
    return paragraph

def add_shading_to_paragraph(paragraph, fill_color):
    """Add background color to paragraph"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    paragraph._element.get_or_add_pPr().append(shading_elm)

def is_korean_content(content):
    """Check if content contains Korean (>10% Korean characters)"""
    korean_chars = sum(1 for char in content if '\uAC00' <= char <= '\uD7A3')
    return korean_chars > len(content) * 0.1

def add_image_with_spacing(document, image_path, figure_num, caption_text):
    """Add image with proper spacing to prevent overlap issues"""
    if not os.path.exists(image_path):
        return figure_num

    # MANDATORY: Blank paragraph BEFORE image
    blank_before = document.add_paragraph()
    blank_before.paragraph_format.space_before = Pt(0)
    blank_before.paragraph_format.space_after = Pt(0)

    # Image paragraph with proper spacing
    image_para = document.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_before = Pt(6)
    image_para.paragraph_format.space_after = Pt(3)
    run = image_para.add_run()
    run.add_picture(image_path, width=Inches(5.5))

    # Caption with no gap between image and caption
    caption_para = document.add_paragraph(caption_text)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_before = Pt(0)
    caption_para.paragraph_format.space_after = Pt(3)
    caption_para.runs[0].font.size = Pt(9)
    caption_para.runs[0].font.italic = True
    caption_para.runs[0].font.color.rgb = RGBColor(127, 140, 141)

    # MANDATORY: Blank paragraph AFTER caption
    blank_after = document.add_paragraph()
    blank_after.paragraph_format.space_before = Pt(0)
    blank_after.paragraph_format.space_after = Pt(0)

    return figure_num + 1

def create_document_with_korean_font():
    """Create new DOCX document with Korean font and tight spacing"""
    document = Document()

    # Set default Korean font and tight spacing
    style = document.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(14)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    return document

print("✅ DOCX setup complete - all helper functions ready")
